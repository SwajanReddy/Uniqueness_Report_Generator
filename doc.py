from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn

def format_table_in_word(doc, table_text):
    """
    Convert Markdown table text into a Word table and add it to the document.
    
    :param doc: The Word document object.
    :param table_text: The Markdown table text to convert.
    """
    # Split the Markdown table into rows
    lines = table_text.strip().split('\n')
    
    if len(lines) < 3:
        print("Not enough lines to form a table.")
        return
    
    # Extract headers and rows
    headers = [header.strip() for header in lines[1].strip().split('|')[1:-1]]
    rows = [line.strip().split('|')[1:-1] for line in lines[2:]]

    if not headers:
        print("No headers found in table.")
        return
    headers = ["Category",	"Details"]
    # Add table to the document
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Add header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
    
    # Add data rows
    for row in rows:
        if len(row) != len(headers):
            print(f"Row length {len(row)} does not match header length {len(headers)}.")
            continue
        cells = table.add_row().cells
        for i, cell in enumerate(row):
            cells[i].text = cell.strip()
    
    # Optional: Adjust column width
    for col in table.columns:
        for cell in col.cells:
            cell.width = Inches(1.5)

def create_word_document(data, file_name='output.docx'):
    """
    Create a Word document from given data.
    
    :param data: Dictionary containing the table, summary, and challenges.
    :param file_name: The name of the file to save the document as.
    """
    # Create a new Document
    doc = Document()
    
    # Add a title or heading
    doc.add_heading('Uniqueness Report', level=1)
    
    # Add formatted table
    if 'table' in data:
        doc.add_heading('Table Summary', level=2)
        format_table_in_word(doc, data['table'])
    
    # Add summary
    if 'summary' in data:
        doc.add_heading('Summary', level=2)
        doc.add_paragraph(data['summary'])
    
    # Add challenges
    if 'challenges' in data:
        doc.add_heading('Challenges to grow', level=2)
        doc.add_paragraph(data['challenges'])
    
    # Save the document
    doc.save(file_name)
    return doc
